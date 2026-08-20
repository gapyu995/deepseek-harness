# -*- coding: utf-8 -*-
"""
PDF -> Word (.docx) extraction with layout preservation.

Extracts text, tables and images from a structured PDF and rebuilds a Word
document that preserves headings, paragraph first-line indentation, list
hanging indentation, and clause numbering. Tuned for Chinese standards
(GB/T) but works on any PDF with regular margins and numbered clauses.

Usage:
    python pdf_to_word.py input.pdf [output.docx]

Dependencies (see requirements.txt): pdfplumber, python-docx, Pillow

Pipeline (see SKILL.md for the full method):
    1. char-level line grouping (not extract_text, to keep mixed CJK/Latin)
    2. measure page margin (odd/even fallback) and indent offsets
    3. classify lines by content + position (chapter/clause/term/body/list/note/ref)
    4. merge continuation lines into paragraphs
    5. extract tables (find_tables) and exclude them from the text stream
    6. decode embedded images natively and place them by bbox order
    7. emit Word with first-line / hanging / left indents in points
"""

import io
import re
import sys
from pathlib import Path

import pdfplumber
from PIL import Image
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
if len(sys.argv) < 2:
    sys.stderr.write('usage: python pdf_to_word.py input.pdf [output.docx]\n')
    sys.exit(2)

SRC = Path(sys.argv[1]).resolve()
if not SRC.exists():
    sys.stderr.write(f'input not found: {SRC}\n')
    sys.exit(2)

OUT = Path(sys.argv[2]).resolve() if len(sys.argv) >= 3 else SRC.with_suffix('.docx')

# ---------------------------------------------------------------------------
# 1. image extraction (decode embedded Image XObjects via pdfplumber streams)
# ---------------------------------------------------------------------------
def _raw_pixels_to_image(data, im):
    """Decode a pdfplumber image dict's stream data into a PIL Image."""
    # JPEG (DCTDecode) / JPEG2000 (JPXDecode) data is a self-contained file.
    if data[:2] == b'\xff\xd8' or data[:12] == b'\x00\x00\x00\x0cjP  \r\n\x87\n':
        img = Image.open(io.BytesIO(data))
        img.load()
        return img
    w, h = im.get('srcsize') or (int(im['width']), int(im['height']))
    bits = im.get('bits', 8)
    cs = im.get('colorspace')
    if bits == 8:
        mode = 'RGB' if cs and 'RGB' in str(cs) else 'L'
        return Image.frombytes(mode, (w, h), data)
    if bits == 1:
        rb = (w + 7) // 8
        img = Image.new('L', (w, h))
        px = img.load()
        for y in range(h):
            base = y * rb
            for x in range(w):
                idx = base + (x >> 3)
                if idx >= len(data):
                    break
                bit = (data[idx] >> (7 - (x & 7))) & 1
                px[x, y] = 255 if bit else 0
        return img
    return None


def materialize(im):
    """Return a PNG BytesIO for a pdfplumber image dict, or None."""
    stream = im.get('stream')
    if stream is None:
        return None
    try:
        data = stream.get_data()
    except Exception:
        return None
    try:
        img = _raw_pixels_to_image(data, im)
    except Exception:
        return None
    if img is None:
        return None
    img = img.convert('RGB')
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# 2. font/style helpers
# ---------------------------------------------------------------------------
def set_run(run, size=10.5, bold=False, ea='宋体'):
    if ea is None:
        ea = '宋体'
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rF = rPr.get_or_add_rFonts()
    rF.set(qn('w:eastAsia'), ea)
    rF.set(qn('w:ascii'), 'Times New Roman')
    rF.set(qn('w:hAnsi'), 'Times New Roman')


def para(doc, text='', size=10.5, bold=False, ea='宋体', align=None,
         before=0, after=2, left=None, first=None, hanging=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = 1.3
    if align is not None:
        p.alignment = align
    if left is not None:
        pf.left_indent = Pt(left)
    if first is not None:
        pf.first_line_indent = Pt(first)
    if hanging is not None:
        pf.left_indent = Pt(hanging[0])
        pf.first_line_indent = Pt(-hanging[1])
    if text:
        run = p.add_run(text)
        set_run(run, size, bold, ea)
    return p


# ---------------------------------------------------------------------------
# 3. collect tables, grouping cross-page continuations by header
# ---------------------------------------------------------------------------
def _clean_row(row):
    return [(c or '').replace('\n', ' ').strip() for c in row]


def _looks_like_heading_cell(cells):
    # single-cell pseudo-tables such as a centered "附录 X" heading
    return len(cells) == 1 and re.match(r'^(附录|附\s*录)', cells[0])


segments = []
with pdfplumber.open(str(SRC)) as pdf:
    for pageno, page in enumerate(pdf.pages, 1):
        for t in page.find_tables():
            segments.append((pageno, t.bbox[1], t))
segments.sort(key=lambda s: (s[0], s[1]))

table_groups = []      # [{header, rows, note}]
seg_group = {}         # (pageno, top) -> group index

for pageno, top, t in segments:
    rows = [_clean_row(r) for r in t.extract()]
    rows = [r for r in rows if not all(c == '' for c in r)]
    if not rows:
        continue
    header = rows[0]
    if _looks_like_heading_cell(header):
        continue
    if table_groups and tuple(header) == tuple(table_groups[-1]['header']):
        gi = len(table_groups) - 1
        table_groups[-1]['rows'].extend(rows[1:])
    else:
        data, note = [], None
        for r in rows[1:]:
            if r[0].startswith('注'):
                note = r[0]
            else:
                data.append(r)
        table_groups.append({'header': header, 'rows': data, 'note': note})
        gi = len(table_groups) - 1
    seg_group[(pageno, top)] = gi


# ---------------------------------------------------------------------------
# 4. classification
# ---------------------------------------------------------------------------
RE_H1 = re.compile(r'^[1-9][\s\u3000]+|^(前|引)[\s\u3000]*(言)|^参考文献|^附[\s\u3000]*录[\s\u3000]*[A-Z]')
RE_CLAUSE = re.compile(r'^((?:[A-Z]|\d+)(?:\.\d+)+)[\s\u3000]+')
RE_TERMNUM = re.compile(r'^((?:[A-Z]|\d+)(?:\.\d+)+)$')
RE_LIST = re.compile(r'^(?:[a-z]|\d+)[）)]')
RE_NOTE = re.compile(r'^(注|示例)[:：]')
RE_SOURCE = re.compile(r'^\[来源')
RE_REF = re.compile(r'^GB/?T\s*\d[\d.]*(?:—\d{4})?[\s\u3000]|^[\[［]\d+[\]］]')
RE_CAPTION = re.compile(r'^(图|表)\s*[0-9A-Z一二三四五六七八九十]')
RE_TOC = re.compile(r'\.{6,}')


def classify(it, prev_type, page_width):
    t = it['text']
    font = it.get('font', '')
    is_heading_font = 'Hei' in font  # SimHei 黑体 marks headings/terms
    centered = abs((it['x0'] + it['x1']) / 2 - page_width / 2) < 60 and (it['x1'] - it['x0']) < page_width * 0.7
    if centered:
        if RE_CAPTION.match(t):
            return 'caption'
        if it['size'] >= 13:
            return 'title'
        return 'formula'
    if RE_TOC.search(t):
        return 'toc'
    if is_heading_font:
        if RE_H1.match(t):
            return 'h1'
        if RE_CLAUSE.match(t):
            return 'h2'
        return 'h2'
    if it['size'] >= 13:
        return 'title'
    off = it['x0'] - it['margin']
    if off < 5:
        if RE_H1.match(t):
            return 'h1'
        m = RE_CLAUSE.match(t)
        if m:
            body = t[m.end():]
            if body.endswith('。') or len(body) > 24:
                return 'clause'
            return 'h2'
        if RE_TERMNUM.match(t):
            return 'term_num'
        return 'cont'
    if off < 25:
        if RE_NOTE.match(t):
            return 'note'
        if RE_LIST.match(t):
            return 'list'
        if RE_SOURCE.match(t):
            return 'source'
        if RE_REF.match(t):
            return 'ref'
        if prev_type == 'term_num':
            return 'term_title'
        return 'body_first'
    return 'cont_indent'


def page_margin(page, pageno, parity_margin):
    xs = [c['x0'] for c in page.chars if 90 <= c['top'] <= 775 and c['x0'] < 400]
    if not xs:
        return parity_margin.get(pageno % 2, 70.0)
    m = min(xs)
    if m > 85 and parity_margin:
        return parity_margin.get(pageno % 2, m)
    return m


# ---------------------------------------------------------------------------
# 5. extract lines + classify + merge
# ---------------------------------------------------------------------------
cover_lines = []
cover_images = []
all_items = []

with pdfplumber.open(str(SRC)) as pdf:
    page_widths = [p.width for p in pdf.pages]
    _m = {}
    for pageno, page in enumerate(pdf.pages, 1):
        xs = [c['x0'] for c in page.chars if 90 <= c['top'] <= 775 and c['x0'] < 400]
        if xs and min(xs) <= 85:
            _m.setdefault(pageno % 2, []).append(min(xs))
    parity_margin = {k: sorted(v)[len(v) // 2] for k, v in _m.items()}

    for pageno, page in enumerate(pdf.pages, 1):
        pw = page.width
        M = page_margin(page, pageno, parity_margin)
        tables_this_page = page.find_tables()
        table_bboxes = [t.bbox for t in tables_this_page]
        chars = sorted(page.chars, key=lambda c: (round(c['top'], 2), c['x0']))
        groups = []
        for c in chars:
            if groups and abs(c['top'] - groups[-1][-1]['top']) < 5.0:
                groups[-1].append(c)
            else:
                groups.append([c])
        for g in groups:
            g = sorted(g, key=lambda c: c['x0'])
            # Drop ASCII-space glyphs that overlap the previous glyph — the PDF
            # generator emits them as positioning artifacts inside CJK glyphs.
            # Full-width space (U+3000) is a meaningful separator and is kept.
            parts = []
            prev_x1 = None
            for c in g:
                if c['text'] == ' ':
                    # drop only spaces that overlap the previous glyph by >1pt
                    if prev_x1 is not None and c['x0'] < prev_x1 - 1.0:
                        continue
                    parts.append(c['text'])
                else:
                    parts.append(c['text'])
                    prev_x1 = c['x1']
            text = ''.join(parts).strip()
            if not text:
                continue
            top = min(c['top'] for c in g)
            x0 = min(c['x0'] for c in g)
            x1 = max(c['x1'] for c in g)
            size = max(c['size'] for c in g)
            fc = {}
            for c in g:
                fn = c.get('fontname', '')
                fc[fn] = fc.get(fn, 0) + 1
            font = max(fc, key=fc.get) if fc else ''
            if top < 90 and (x0 > pw * 0.6 or re.match(r'^GB', text)):
                continue
            if top > 775:
                continue
            if any(y0 - 3 <= top <= y1 + 3 for (_x0, y0, _x1, y1) in table_bboxes):
                continue
            if pageno == 1:
                cover_lines.append((top, text, size))
            else:
                all_items.append((pageno, top, 'text',
                                  dict(text=text, x0=x0, x1=x1, size=size, margin=M, font=font)))
        for t in tables_this_page:
            gi = seg_group.get((pageno, t.bbox[1]))
            if gi is not None:
                all_items.append((pageno, t.bbox[1], 'table', gi))
        for im in page.images:
            if pageno == 1:
                cover_images.append(im)
            else:
                all_items.append((pageno, im['top'], 'image', im))

all_items.sort(key=lambda it: (it[0], it[1]))
cover_lines.sort(key=lambda x: x[0])

blocks = []


def last_type():
    return blocks[-1]['type'] if blocks else None


for (pg, _top, kind, payload) in all_items:
    if kind == 'table':
        blocks.append({'type': 'table', 'gi': payload, 'pg': pg})
        continue
    if kind == 'image':
        blocks.append({'type': 'image', 'im': payload, 'pg': pg})
        continue
    role = classify(payload, last_type(), page_widths[pg - 1])
    t = payload['text']
    size = payload.get('size', 10.5)
    base = {'text': t, 'pg': pg, 'size': size}
    if role == 'title':
        blocks.append({'type': 'title', **base})
    elif role == 'formula':
        blocks.append({'type': 'formula', **base})
    elif role == 'toc':
        blocks.append({'type': 'toc', **base})
    elif role == 'caption':
        blocks.append({'type': 'caption', **base})
    elif role == 'h1':
        blocks.append({'type': 'h1', **base})
    elif role == 'h2':
        blocks.append({'type': 'h2', **base})
    elif role == 'clause':
        blocks.append({'type': 'clause', **base})
    elif role == 'term_num':
        blocks.append({'type': 'term_num', **base})
    elif role == 'term_title':
        if last_type() == 'term_num':
            blocks[-1]['type'] = 'term'
            blocks[-1]['text'] += '\u3000' + t
        else:
            blocks.append({'type': 'para', **base})
    elif role == 'body_first':
        blocks.append({'type': 'para', **base})
    elif role == 'list':
        blocks.append({'type': 'list', **base})
    elif role == 'note':
        blocks.append({'type': 'note', **base})
    elif role == 'source':
        blocks.append({'type': 'source', **base})
    elif role == 'ref':
        blocks.append({'type': 'ref', **base})
    elif role in ('cont', 'cont_indent'):
        if blocks and blocks[-1].get('text') is not None:
            blocks[-1]['text'] += t
        else:
            blocks.append({'type': 'para', **base})


# ---------------------------------------------------------------------------
# 6. emit Word
# ---------------------------------------------------------------------------
doc = Document()
st = doc.styles['Normal']
st.font.name = 'Times New Roman'
st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def emit_term(b):
    t = b['text']
    m = re.match(r'^(\d+(?:\.\d+)+)\u3000(.*)$', t)
    if not m:
        para(doc, t, 10.5, True, '黑体', before=6)
        return
    num, rest = m.group(1), m.group(2)
    parts = rest.split('\u3000')
    para(doc, '', before=6, after=1)
    pr = doc.paragraphs[-1]
    r = pr.add_run(num + '\u3000')
    set_run(r, 10.5, False, '宋体')
    r = pr.add_run(parts[0])
    set_run(r, 10.5, True, '黑体')
    if len(parts) > 1:
        r = pr.add_run('\u3000' + '\u3000'.join(parts[1:]))
        set_run(r, 10.5, False, '宋体')


# cover (page 1)
for im in cover_images:
    buf = materialize(im)
    if buf is not None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(buf, width=Cm(4.0))
for _top, text, size in cover_lines:
    if size >= 20:
        para(doc, text, 22, True, '黑体', WD_ALIGN_PARAGRAPH.CENTER, before=2, after=4)
    elif size >= 14:
        para(doc, text, 16, True, '黑体', WD_ALIGN_PARAGRAPH.CENTER, before=2, after=4)
    elif size >= 11:
        para(doc, text, 13, False, '宋体', WD_ALIGN_PARAGRAPH.CENTER, before=2, after=4)
    else:
        para(doc, text, 9, False, '宋体', WD_ALIGN_PARAGRAPH.CENTER, before=0, after=2)
doc.add_page_break()

def emit_table(g):
    ncols = len(g['header'])
    tbl = doc.add_table(rows=1, cols=ncols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htxt in enumerate(g['header']):
        c = tbl.rows[0].cells[i]
        c.text = ''
        r = c.paragraphs[0].add_run(htxt)
        set_run(r, 9, True, '黑体')
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in g['rows']:
        cells = tbl.add_row().cells
        for i in range(ncols):
            cells[i].text = ''
            r = cells[i].paragraphs[0].add_run(row[i] if i < len(row) else '')
            set_run(r, 9, False, '宋体')
            if i != 1:
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if g.get('note'):
        para(doc, g['note'], 9, False, '宋体', after=2)


emitted_groups = set()


def sz(b, dflt=10.5):
    s = b.get('size')
    if not s:
        return dflt
    return round(s * 2) / 2  # round to 0.5pt, preserving the source size


last_pg = blocks[0]['pg'] if blocks else None
for b in blocks:
    if b.get('pg') is not None and last_pg is not None and b['pg'] != last_pg:
        doc.add_page_break()
        last_pg = b['pg']
    ty = b['type']
    s = sz(b)
    if ty == 'title':
        para(doc, b['text'], s, True, '黑体', WD_ALIGN_PARAGRAPH.CENTER, before=4, after=4)
    elif ty == 'formula':
        para(doc, b['text'], s, False, '宋体', WD_ALIGN_PARAGRAPH.CENTER, after=2)
    elif ty == 'caption':
        if '续' in b['text']:
            continue
        para(doc, b['text'], s, True, '宋体', WD_ALIGN_PARAGRAPH.CENTER, before=3, after=3)
    elif ty == 'h1':
        para(doc, b['text'], s, True, '黑体', before=6, after=3)
    elif ty == 'h2':
        para(doc, b['text'], s, True, '黑体', before=4, after=2)
    elif ty == 'clause':
        para(doc, b['text'], s, False, '宋体', after=2)
    elif ty == 'para':
        para(doc, b['text'], s, False, '宋体', after=2, first=21)
    elif ty == 'term':
        emit_term(b)
    elif ty == 'list':
        para(doc, b['text'], s, False, '宋体', after=1, hanging=(42, 21))
    elif ty == 'note':
        para(doc, b['text'], s, False, '宋体', after=1, hanging=(34, 16))
    elif ty == 'source':
        para(doc, b['text'], s, False, '宋体', after=2, left=21)
    elif ty == 'ref':
        para(doc, b['text'], s, False, '宋体', after=2, left=21)
    elif ty == 'toc':
        para(doc, b['text'], s, False, '宋体', after=1)
    elif ty == 'image':
        im = b['im']
        if im.get('srcsize') and im['srcsize'][1] < 60:
            continue  # skip inline glyph images
        buf = materialize(im)
        if buf is not None:
            width_cm = min(max((im['x1'] - im['x0']) * 0.0353, 0.3), 16.0)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(buf, width=Cm(width_cm))
    elif ty == 'table':
        gi = b['gi']
        if gi in emitted_groups:
            continue
        emitted_groups.add(gi)
        emit_table(table_groups[gi])

doc.save(str(OUT))
print(f'saved {OUT}')
