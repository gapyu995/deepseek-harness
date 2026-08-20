# -*- coding: utf-8 -*-
"""
PDF -> Word with identical page layout (image mode).

Renders every PDF page to a high-resolution image and embeds each as a
full-page picture in a Word document. The .docx reproduces the source page for
page and pixel for pixel; text is not selectable because each page is an image.

Usage:
    python pdf_to_word_layout.py input.pdf [output.docx] [dpi]

Dependencies: pypdfium2, python-docx, Pillow
"""

import io
import sys
from pathlib import Path

import pypdfium2 as pdfium
from docx import Document
from docx.shared import Cm

A4_W_CM = 21.0
A4_H_CM = 29.7


def main():
    if len(sys.argv) < 2:
        sys.stderr.write('usage: python pdf_to_word_layout.py input.pdf [output.docx] [dpi]\n')
        return 2
    src = Path(sys.argv[1]).resolve()
    if not src.exists():
        sys.stderr.write(f'input not found: {src}\n')
        return 2
    out = Path(sys.argv[2]).resolve() if len(sys.argv) >= 3 else src.with_suffix('.docx')
    dpi = int(sys.argv[3]) if len(sys.argv) >= 4 else 200
    scale = dpi / 72.0

    pdf = pdfium.PdfDocument(str(src))
    n = len(pdf)

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(A4_W_CM)
    sec.page_height = Cm(A4_H_CM)
    sec.top_margin = Cm(0)
    sec.bottom_margin = Cm(0)
    sec.left_margin = Cm(0)
    sec.right_margin = Cm(0)

    for i in range(n):
        bmp = pdf[i].render(scale=scale)
        pil = bmp.to_pil()
        buf = io.BytesIO()
        pil.save(buf, format='PNG')
        buf.seek(0)

        p = doc.add_paragraph()
        p.paragraph_format.space_before = 0
        p.paragraph_format.space_after = 0
        p.paragraph_format.line_spacing = 1.0
        # fit width; the page image aspect matches A4 so height follows
        p.add_run().add_picture(buf, width=Cm(A4_W_CM))
        if i < n - 1:
            doc.add_page_break()

    doc.save(str(out))
    print(f'saved {out} ({n} pages at {dpi} dpi)')
    return 0


if __name__ == '__main__':
    sys.exit(main())
